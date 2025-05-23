import os
import docx
import fitz
import pandas as pd

def extraer_texto_de_pdf(file_path):
    text = ""
    with fitz.open(file_path) as doc:
        for page in doc:
            text += page.get_text("text")
    return text

def extraer_texto_de_word(file_path):
    doc = docx.Document(file_path)
    text = []
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                text.append(row_text)
    return "\n".join(text)

def extraer_texto_de_excel(file_path):
    dfs = pd.read_excel(file_path, sheet_name=None)
    text = []
    for sheet_name, df in dfs.items():
        text.append(f"--- Hoja: {sheet_name} ---")
        text.append(df.to_string(index=False, header=True))
    return "\n".join(text)

def main(input_folder, output_file):
    all_text = []

    for file in os.listdir(input_folder):
        path = os.path.join(input_folder, file)
        try:
            text = ""
            if file.lower().endswith(".pdf"):
                text = extraer_texto_de_pdf(path)
            elif file.lower().endswith(".docx"):
                text = extraer_texto_de_word(path)
            elif file.lower().endswith(".xlsx"):
                text = extraer_texto_de_excel(path)
            else:
                continue

            if text.strip():
                header = f"\n\n===== Archivo: {file} =====\n\n"
                all_text.append(header + text)
                print(f"Texto extraído de {file} (primeros 10000 caracteres):")
                print(text[:10000] + "\n" + "-"*30)
            else:
                print(f"No se extrajo texto de {file}")

        except Exception as e:
            print(f"Error con {file}: {e}")

    # guarda todo en el mismo archivo (borrar si no es necesario)
    with open(output_file, "w", encoding="utf-8") as f:
        f.write("\n".join(all_text))

    print(f"\nTexto combinado guardado en {output_file}")

if __name__ == "__main__":
    input_folder = "carpeta_de_archivos"  # Carpeta donde estan subidos los archivos
    output_file = "texto_combinado.txt"   # archivo donde se guarda todo
    main(input_folder, output_file)
